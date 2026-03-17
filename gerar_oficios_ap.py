import os
import pandas as pd
from docx import Document
from docx.shared import Pt  # para fixar Calibri 12 pt

# ALTERAR DE ACORDO COM A OCASIÃO
TEMPLATE = "modelo_oficio_ap_representante.docx"
PLANILHA = "dados_oficios_req1_2026.xlsx"
# PLANILHA = "U:/COCPSF/SACIFR/Ano_CI_2025/Audiências_Públicas/dados_oficios_req80.xlsx"
PASTA_OUT = "oficios_gerados"

# --------------------------------------------------------------------------- #
# 1) Ler a planilha
df = pd.read_excel(PLANILHA)

# Numeração automática se a coluna "n" estiver vazia
if df["n"].isna().all():
    N0 = 1
    df["n"] = range(N0, N0 + len(df))

# Detectar se a coluna do mês tem acento ou não
mes_col = "mês" if "mês" in df.columns else "mes"

# --------------------------------------------------------------------------- #
# 2) Funções auxiliares
# def tratamento_e_pronome(sexo: str):
#     """Retorna ('Ao Senhor', 'Senhor') ou ('À Senhora', 'Senhora')."""
#     sexo = (sexo or "").strip().upper()
#     if sexo.startswith("F"):
#         return "À Senhora", "Senhora"
#     return "Ao Senhor", "Senhor"


def tratamento(sexo: str, cargo: str):
    sexo = (sexo or "").strip().upper()
    cargo = (cargo or "").strip().upper()
    cargoExcelencia = cargo.startswith("MINISTR") or cargo.startswith("COMANDANTE")
    if sexo.startswith("F"):
        if cargoExcelencia:
            return "A Sua Excelência a Senhora"
        return "À Senhora"
    if cargoExcelencia:
        return "A Sua Excelência o Senhor"
    return "Ao Senhor"


def vocativo(sexo, cargo_resumido):
    cargo_resumido_normalized = (cargo_resumido or "").strip().upper()
    if cargo_resumido_normalized == "SENHOR":
        return "Prezado Senhor"
    if cargo_resumido_normalized == "SENHORA":
        return "Prezada Senhora"
    sexo = (sexo or "").strip().upper()
    if sexo.startswith("F"):
        return f"Senhora {cargo_resumido}"
    if sexo.startswith("M"):
        return f"Senhor {cargo_resumido}"


def obj_pronome(sexo: str):
    """Retorna 'o' para masculino, 'a' para feminino."""
    sexo = (sexo or "").strip().upper()
    return "a" if sexo.startswith("F") else "o"


def substituir_paragrafos(paragraphs, mapa):
    """
    Substitui marcadores em cada parágrafo preservando o estilo original
    (bold, itálico, sublinhado etc.), mas forçando fonte Calibri 12 pt.
    """
    for p in paragraphs:
        texto_original = p.text
        texto_novo = texto_original
        for alvo, novo in mapa.items():
            texto_novo = texto_novo.replace(alvo, str(novo))

        if texto_novo != texto_original:
            # Captura o estilo do primeiro run (se existir)
            if p.runs:
                primeiro_run = p.runs[0]
                estilo_base = {
                    "bold": primeiro_run.bold,
                    "italic": primeiro_run.italic,
                    "underline": primeiro_run.underline,
                    "style": primeiro_run.style,
                }
            else:
                estilo_base = {
                    "bold": None,
                    "italic": None,
                    "underline": None,
                    "style": None,
                }

            # Remove todos os runs
            for idx in range(len(p.runs) - 1, -1, -1):
                p.runs[idx].clear()

            # Cria um único run com o texto novo
            run = p.add_run(texto_novo)

            # Aplica Calibri 12 pt
            run.font.name = "Calibri"
            run.font.size = Pt(12)

            # Restaura atributos de estilo base
            run.bold = estilo_base["bold"]
            run.italic = estilo_base["italic"]
            run.underline = estilo_base["underline"]
            run.style = estilo_base["style"]


# --------------------------------------------------------------------------- #
# 3) Loop principal
os.makedirs(PASTA_OUT, exist_ok=True)

for _, linha in df.iterrows():
    if str(linha["oficio_enviado"]).lower() == "s":
        continue
    doc = Document(TEMPLATE)

    # trat, pron = tratamento_e_pronome(linha["sexo"])

    cargo_cap = str(linha["cargo"]).lstrip()  # Maiúscula inicial
    cargo_upper = cargo_cap.upper()  # CAIXA ALTA
    expositor = "expositora" if linha["sexo"] == "F" else "expositor"

    # Mapa de marcadores → valores
    mapa = {
        "[n]": int(linha["n"]),
        "[dia]": linha["dia"],
        "[mês]": linha[mes_col],
        "[Tratamento]": tratamento(linha["sexo"], linha["cargo"]),
        "[vocativo]": vocativo(linha["sexo"], linha["cargo_resumido"]),
        # "[Pronome]":    pron,
        "[objPron]": obj_pronome(linha["sexo"]),
        "[NOME]": linha["nome"],
        "[Cargo]": linha["cargo"],
        "[cargo_resumido]": linha["cargo_resumido"],
        "[entidade]": linha["entidade"],
        # "[entidade_abreviado]": linha["entidade_abreviado"],
        "[entidadePreposicao]": linha["entidadePreposicao"],
        "[expositor]": expositor,
        # ALTERAR OS CAMPOS ABAIXO DE ACORDO COM A REUNIÃO A SER REALIZADA
        "[objetivo]": "Debater acerca da regulamentação e das providências adotadas após o advento da Lei nº 15.097/2025, que trata do aproveitamento de potencial energético offshore'",
        "[requerimentos]": "ao Requerimento nº 1/2026-CI",
        "[data_reuniao]": "7 de abril de 2026, terça-feira",
        "[horario_reuniao]": "09h00",
        "[local_reuniao]": "no Plenário nº 13 da Ala Alexandre Costa, Anexo II, do Senado Federal",
    }

    # Substituição em parágrafos normais
    substituir_paragrafos(doc.paragraphs, mapa)
    # … e em cada célula de cada tabela
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                substituir_paragrafos(cell.paragraphs, mapa)

    # Salvar
    primeiro_nome = linha["nome"].split()[0]
    # ALTERAR NOME PADRÃO DO DOCUMENTO DE ACORDO COM A REUNIÃO A SER REALIZADA
    destino = f"{int(linha['n']):03d} - REQ 1_2026 - {linha['entidade']}.docx"
    doc.save(os.path.join(PASTA_OUT, destino))

print(f"{len(df)} ofício(s) gerado(s) em '{PASTA_OUT}'.")
