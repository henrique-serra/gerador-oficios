import os
import pandas as pd
from docx import Document
from docx.shared import Pt           # para fixar Calibri 12 pt

TEMPLATE   = "modelo_oficio.docx"
PLANILHA   = "dados_oficios.xlsx"
PASTA_OUT  = "oficios_gerados"

# --------------------------------------------------------------------------- #
# 1) Ler a planilha
df = pd.read_excel(PLANILHA)

# Numeração automática se a coluna "n" estiver vazia
if df["n"].isna().all():
    N0 = 1
    df["n"] = range(N0, N0 + len(df))

# Detectar se a coluna do mês tem acento ou não
mes_col = "mês" if "mês" in df.columns else "mes"

# --------------------------------------------------------------------------- #
# 2) Funções auxiliares
def tratamento_e_pronome(sexo: str):
    """Retorna ('Ao Senhor', 'Senhor') ou ('À Senhora', 'Senhora')."""
    sexo = (sexo or "").strip().upper()
    if sexo.startswith("F"):
        return "À Senhora", "Senhora"
    return "Ao Senhor", "Senhor"

def substituir_paragrafos(paragraphs, mapa):
    """
    Substitui marcadores em cada parágrafo preservando o estilo original
    (bold, itálico, sublinhado etc.), mas forçando fonte Calibri 12 pt.
    """
    for p in paragraphs:
        texto_original = p.text
        texto_novo = texto_original
        for alvo, novo in mapa.items():
            texto_novo = texto_novo.replace(alvo, str(novo))

        if texto_novo != texto_original:
            # Captura o estilo do primeiro run (se existir)
            if p.runs:
                primeiro_run = p.runs[0]
                estilo_base  = {
                    "bold":      primeiro_run.bold,
                    "italic":    primeiro_run.italic,
                    "underline": primeiro_run.underline,
                    "style":     primeiro_run.style,
                }
            else:
                estilo_base = {"bold": None, "italic": None,
                               "underline": None, "style": None}

            # Remove todos os runs
            for idx in range(len(p.runs) - 1, -1, -1):
                p.runs[idx].clear()

            # Cria um único run com o texto novo
            run = p.add_run(texto_novo)

            # Aplica Calibri 12 pt
            run.font.name = "Calibri"
            run.font.size = Pt(12)

            # Restaura atributos de estilo base
            run.bold      = estilo_base["bold"]
            run.italic    = estilo_base["italic"]
            run.underline = estilo_base["underline"]
            run.style     = estilo_base["style"]

# --------------------------------------------------------------------------- #
# 3) Loop principal
os.makedirs(PASTA_OUT, exist_ok=True)

for _, linha in df.iterrows():
    doc = Document(TEMPLATE)

    trat, pron = tratamento_e_pronome(linha["sexo"])

    cargo_cap   = str(linha["cargo"]).lstrip()     # Maiúscula inicial
    cargo_upper = cargo_cap.upper()                # CAIXA ALTA

    # Mapa de marcadores → valores
    mapa = {
        "[n]":          int(linha["n"]),
        "[dia]":        linha["dia"],
        "[mês]":        linha[mes_col],
        "[Tratamento]": trat,
        "[Pronome]":    pron,
        "[NOME]":       linha["nome"],
        "[Cargo]":      cargo_cap,
        "[CARGO]":      cargo_upper,
    }

    # Substituição em parágrafos normais
    substituir_paragrafos(doc.paragraphs, mapa)
    # … e em cada célula de cada tabela
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                substituir_paragrafos(cell.paragraphs, mapa)

    # Salvar
    primeiro_nome = linha["nome"].split()[0]
    destino = f"{int(linha['n']):03d}_{primeiro_nome}.docx"
    doc.save(os.path.join(PASTA_OUT, destino))

print(f"{len(df)} ofício(s) gerado(s) em '{PASTA_OUT}'.")
