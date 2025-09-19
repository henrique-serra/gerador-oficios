import os
import pandas as pd
from docx import Document
from docx.shared import Pt           # para fixar Calibri 12 pt

# ALTERAR DE ACORDO COM A OCASIÃO
TEMPLATE   = "modelo_oficio_informacoes.docx"
PLANILHA   = "dados_oficios_informacoes.xlsx"
PASTA_OUT  = "oficios_gerados"

# --------------------------------------------------------------------------- #
# 1) Ler a planilha
df = pd.read_excel(PLANILHA)

# Numeração automática se a coluna "n" estiver vazia
if df["n_oficio"].isna().all():
    N0 = 1
    df["n_oficio"] = range(N0, N0 + len(df))

# Detectar se a coluna do mês tem acento ou não
mes_col = "mês" if "mês" in df.columns else "mes"

# --------------------------------------------------------------------------- #
# 2) Funções auxiliares
def tratamento_e_pronome(sexo: str):
    """Retorna ('Ao Senhor', 'Senhor') ou ('À Senhora', 'Senhora')."""
    sexo = (sexo or "").strip().upper()
    if sexo.startswith("F"):
        return "À Senhora", "Senhora"
    return "Ao Senhor", "Senhor"

def obj_pronome(sexo: str):
    """Retorna 'o' para masculino, 'a' para feminino."""
    sexo = (sexo or "").strip().upper()
    return "a" if sexo.startswith("F") else "o"

def substituir_paragrafos(paragraphs, mapa):
    """
    Substitui marcadores em cada parágrafo preservando o estilo original
    (bold, itálico, sublinhado etc.), mas forçando fonte Calibri 12 pt.
    """
    for p in paragraphs:
        texto_original = p.text
        texto_novo = texto_original
        for alvo, novo in mapa.items():
            texto_novo = texto_novo.replace(alvo, str(novo))

        if texto_novo != texto_original:
            # Captura o estilo do primeiro run (se existir)
            if p.runs:
                primeiro_run = p.runs[0]
                estilo_base  = {
                    "bold":      primeiro_run.bold,
                    "italic":    primeiro_run.italic,
                    "underline": primeiro_run.underline,
                    "style":     primeiro_run.style,
                }
            else:
                estilo_base = {"bold": None, "italic": None,
                               "underline": None, "style": None}

            # Remove todos os runs
            for idx in range(len(p.runs) - 1, -1, -1):
                p.runs[idx].clear()

            # Cria um único run com o texto novo
            run = p.add_run(texto_novo)

            # Aplica Calibri 12 pt
            run.font.name = "Calibri"
            run.font.size = Pt(12)

            # Restaura atributos de estilo base
            run.bold      = estilo_base["bold"]
            run.italic    = estilo_base["italic"]
            run.underline = estilo_base["underline"]
            run.style     = estilo_base["style"]

# --------------------------------------------------------------------------- #
# 3) Loop principal
os.makedirs(PASTA_OUT, exist_ok=True)

for _, linha in df.iterrows():
    doc = Document(TEMPLATE)

    trat, pron = tratamento_e_pronome(linha["sexo"])

    cargo_cap   = str(linha["cargo"]).lstrip()     # Maiúscula inicial
    cargo_upper = cargo_cap.upper()                # CAIXA ALTA

    # Mapa de marcadores → valores
    mapa = {
        "[n_oficio]": int(linha["n_oficio"]),
        "[dia_oficio]": linha["dia_oficio"],
        "[mes_oficio]": linha["mes_oficio"],
        "[ano_oficio]": linha["ano_oficio"],
        "[Tratamento]": trat,
        "[Pronome]":    pron,
        "[objPron]":    obj_pronome(linha["sexo"]),
        "[NOME]":       linha["nome"],
        "[Cargo]":      linha["cargo"],
        "[cargo_resumido]": linha["cargo_resumido"],
        "[entidade]":   linha["entidade"],
        # "[entidade_abreviado]": linha["entidade_abreviado"],
        "[entidadePreposicao]": linha["entidadePreposicao"],
        "[n_reuniao]": linha["n_reuniao"],
        "[extraordinaria]": linha["extraordinaria"],
        "[data_reuniao]": linha["data_reuniao"],
        "[req_num]": linha["req_num"],
        "[req_ano]": linha["req_ano"],
        "[por_parte]": linha["por_parte"],
        "[presidente_comissao]": "Senador Marcos Rogério"
    }

    # Substituição em parágrafos normais
    substituir_paragrafos(doc.paragraphs, mapa)
    # … e em cada célula de cada tabela
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                substituir_paragrafos(cell.paragraphs, mapa)

    # Salvar
    primeiro_nome = linha["nome"].split()[0]
    # ALTERAR NOME PADRÃO DO DOCUMENTO DE ACORDO COM A REUNIÃO A SER REALIZADA
    destino = f"{int(linha['n_oficio']):03d} - REQ {linha['req_num']} - {linha['entidade']}.docx"
    doc.save(os.path.join(PASTA_OUT, destino))

print(f"{len(df)} ofício(s) gerado(s) em '{PASTA_OUT}'.")
